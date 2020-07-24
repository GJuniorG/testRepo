import os
import zipfile
import uuid
import lxml.etree as ET

import docx
import docx2txt
import shutil
import sys
import time
import re
from transformers import MarianMTModel, MarianTokenizer
from docx import Document

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML

start_time = time.time()

p = re.compile('[a-zA-Z1-9]{2,}\s+[a-zA-Z1-9]{2,}')

UPLOAD_FOLDER_INPUT = 'C:/Users/shipant/PycharmProjects/NLP_LanguageTranslation/FileStorage_Input'

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'


def translat(src_text):
    time1 = time.time()
    translated = model.generate(**tokenizer.prepare_translation_batch(src_text))
    tgt_text = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
    time2 = time.time()
    duration = (time2 - time1) * 1000.0
    return tgt_text, duration


def get_para_list(path):
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.iter(PARA):
        texts = [node.text
                 for node in paragraph.iter(TEXT)
                 if node.text]
        if texts:
            paragraphs.append(''.join(texts))
    return paragraph

def get_docx_text(lst):
    global model_name
    global tokenizer
    global model

    model_name = 'Helsinki-NLP/opus-mt-en-de'
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    for str in lst:
        print(str)
        text = str.split("\n")
        text = list(filter(None, text))
        text = [s for s in text if p.match(s)]

        text = [">>de<< " + s for s in text]

        i = 0
        document = Document()
        for textblock in chunks(text, 5):
            i = i + 1
            print("batch #%i (len: %s)" % (i, len(textblock)), file=sys.stderr)
            print("\t " + str(tuple(textblock)))
            target, duration = translat(textblock)
            document.add_paragraph(target)
            print("\t " + str(tuple(target)))
            print('translate took {:.3f} ms'.format(duration), file=sys.stderr)
            print("\n\n")
        end_time = time.time()
    document.save("new.docx")



def chunks(lst, n):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), n):
        yield list(lst[i:i + n])

def translateDocx(source, target, file):
    global model_name
    global tokenizer
    global model

    if source == 'en' and target == 'de':
        model_name = 'Helsinki-NLP/opus-mt-en-de'
    if source == 'de' and target == 'en':
        model_name = 'Helsinki-NLP/opus-mt-de-en'

    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)

    text = docx2txt.process(file).split("\n")
    text = list(filter(None, text))
    text = [s for s in text if p.match(s)]

    text = [">>de<< " + s for s in text]

    i = 0
    document = Document()
    for textblock in chunks(text, 2):
        i = i + 1
        print("batch #%i (len: %s)" % (i, len(textblock)), file=sys.stderr)
        print("\t " + str(tuple(textblock)))
        target, duration = translat(textblock)
        document.add_paragraph(target)
        print("\t " + str(tuple(target)))
        print('translate took {:.3f} ms'.format(duration), file=sys.stderr)
        print("\n\n")
    end_time = time.time()
    duration = (start_time - end_time) * 1000.0
    print('Total translate took {:.3f} ms'.format(duration))
    return document

def replace_string2(filename):
    global model_name
    global tokenizer
    global model

    model_name = 'Helsinki-NLP/opus-mt-en-de'


    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)

    document = zipfile.ZipFile(filename)
    xml_content = document.read('word/document.xml')
    #document.close()
    tree = XML(xml_content)
    # using lxml instead of xml preserved the comments

    paragraphs = []
    i =0
    for paragraph in tree.iter(PARA):
        i=i+1
        texts = [node.text
                 for node in paragraph.iter(TEXT)
                 if node.text]
        if texts:
            #text = list(filter(None, text))
            #text = [s for s in text if p.match(s)]

            #text = [">>de<< " + s for s in text]
            #print("%s: %s" %(i,texts))
            target, duration = translat(texts)
            paragraph.text.replace(texts,target)

    document.save("new.docx")



if __name__ == '__main__':
    #get_para_list("abc.docx")
    replace_string2("abc.docx")
