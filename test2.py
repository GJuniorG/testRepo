import re

import docx
from docx import Document
from transformers import MarianMTModel, MarianTokenizer

p = re.compile('[a-zA-Z1-9]{2,}\s+[a-zA-Z1-9]{2,}')


def trans(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def chunks(lst, n):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), n):
        yield list(lst[i:i + n])


def replace_string2(filename):
    global model_name
    global tokenizer
    global model

    model_name = 'Helsinki-NLP/opus-mt-en-de'
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    doc = Document(filename)
    for p in doc.paragraphs:
        textsL = [p.text]
        textsL = list(filter(None, textsL))

        for text_value in textsL:
            textElem = [text_value]
            translated_as_L = translat(textElem)
            for translatedE in translated_as_L:
                print(text_value)
                print(translatedE)
                text = p.text.replace(text_value, translatedE)
                style = p.style
                p.text = text
                p.style = style

    # doc.save(filename)
    doc.save('test.docx')
    return 1


def translat(src_text):
    translated = model.generate(**tokenizer.prepare_translation_batch(src_text))
    tgt_text = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]

    return tgt_text


if __name__ == '__main__':
    # get_para_list("abc.docx")
    # print(trans("abc.docx"))
    replace_string2("abc.docx")
